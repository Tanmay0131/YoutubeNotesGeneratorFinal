from youtube_transcript_api import YouTubeTranscriptApi
import os
import google.generativeai as genai
from docx import Document
from apiclient import discovery
from httplib2 import Http
from oauth2client import client, file, tools
import webbrowser
from oauth2client.file import Storage
import re
import json
import io
from flask import Flask, request, send_file

def process_text():
    link = request.form.get('youtube_link')
    genai.configure(api_key=request.form.get('api_key'))
    transcript = YouTubeTranscriptApi.get_transcript(link, languages=['en'])
    combined_text = ' '.join([item['text'] for item in transcript])

    prompt = f"""
    Make clean and organized notes using the following transcript. Make sure the information is relevant and concise.

    Transcript: "{combined_text}"
    """
    response = genai.GenerativeModel('gemini-1.5-flash').generate_content(prompt)

    # Save the notes as a Word document
    if hasattr(response, 'text'):
        cleaned_text = response.text.replace('*', '')

        response_doc = io.BytesIO()
        document = Document()
        document.add_heading('Generated Notes', level=1)
        document.add_paragraph('Taken from: https://www.youtube.com/watch?v=' + link)
        document.add_paragraph(cleaned_text)
        document.save(response_doc)
        response_doc.seek(0)

        # current_directory = os.getcwd()
        # response_file_path = os.path.join(current_directory, 'response.docx')
        # document.save(response_file_path)
        print("Response saved to response.docx")

        # Generate questions based on the notes and transcript
        question_prompt = f"""
        Based on the following notes and transcript, generate a list of questions without answers that would help someone study the material.

        Notes: "{cleaned_text}"
        Transcript: "{combined_text}"
        """
        question_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(question_prompt)

        # Save the questions as a Word document
        if hasattr(question_response, 'text'):
            questions_text = question_response.text.replace('*', '')

            questions_doc = io.BytesIO()
            questions_document = Document()
            questions_document.add_heading('Generated Questions', level=1)
            questions_document.add_paragraph('Taken from: https://www.youtube.com/watch?v=' + link)
            questions_document.add_paragraph(questions_text)
            document.save(questions_doc)
            questions_doc.seek(0)

            # questions_file_path = os.path.join(current_directory, 'questions.docx')
            # questions_document.save(questions_file_path)
            print("Questions saved to questions.docx")

            # Generate answers based on the questions, notes, and transcript
            answers_prompt = f"""
            Using only knowledge from the provided documents, provide answers to the following questions based on the notes and transcript provided.

            Questions: "{questions_text}"
            Notes: "{cleaned_text}"
            Transcript: "{combined_text}"
            """
            answers_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(answers_prompt)

            print("Answers saved to answers.docx")

            # Save the answers as a Word document
            if hasattr(answers_response, 'text'):
                answers_text = answers_response.text.replace('*', '')

                answers_doc = io.BytesIO()
                answers_document = Document()
                answers_document.add_heading('Generated Answers', level=1)
                answers_document.add_paragraph('Taken from: https://www.youtube.com/watch?v=' + link)
                answers_document.add_paragraph(answers_text)
                answers_document.save(answers_doc)
                answers_doc.seek(0)

                # answers_file_path = os.path.join(current_directory, 'answers.docx')
                # answers_document.save(answers_file_path)

                # Generate the form prompt asking for JSON format
                form_prompt = f"""
                Based on the following notes and transcript, generate a list of multiple-choice questions that can be used in a quiz. Each question should have four answer choices. Format each question and its answer choices as JSON objects with the following structure:

                {{
                    "question": "Question text",
                    "options": [
                        "Option A",
                        "Option B",
                        "Option C",
                        "Option D"
                    ],
                    "correctAnswer": "A"  # The correct answer as a single letter (A, B, C, or D)
                }}

                Notes: "{cleaned_text}"
                Transcript: "{combined_text}"
                Ensure the JSON is valid and can be parsed directly. Do not use any formatting.
                """

                form_question_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(form_prompt)

                # Assuming `form_question_response` contains the AI-generated JSON text.
                if hasattr(form_question_response, 'text'):
                    forms_questions_json = form_question_response.text.strip()

                    print(forms_questions_json)
                    # Save the raw response to a file for inspection in case of errors
                    with open('response.docx', 'w') as response_file:
                        response_file.write(forms_questions_json)

                    try:
                        # Parse the JSON response
                        forms_questions = json.loads(forms_questions_json)
                        print("Parsed JSON good!")

                        # Save the parsed questions and answers to files for inspection
                        questions = [q.get("question", "") for q in forms_questions]
                        with open('questions.docx', 'w') as questions_file:
                            questions_file.write('\n'.join(questions))

                        answers = [q.get("correctAnswer", "") for q in forms_questions]
                        with open('answers.docx', 'w') as answers_file:
                            answers_file.write('\n'.join(answers))

                        # Setup Google Forms API
                        SCOPES = [
                            "https://www.googleapis.com/auth/forms.body",
                            "https://www.googleapis.com/auth/forms.responses.readonly"
                        ]
                        DISCOVERY_DOC = "https://forms.googleapis.com/$discovery/rest?version=v1"

                        # Setup the authentication flow
                        store = Storage("token.json")
                        creds = store.get()
                        if not creds or creds.invalid:
                            flow = client.flow_from_clientsecrets("client_secrets.json", SCOPES)
                            creds = tools.run_flow(flow, store)

                        # Build the form service
                        form_service = discovery.build(
                            "forms",
                            "v1",
                            http=creds.authorize(Http()),
                            discoveryServiceUrl=DISCOVERY_DOC,
                            static_discovery=False,
                        )

                        # Define the initial form structure
                        form = {
                            "info": {
                                "title": "Generated Quiz",
                            }
                        }

                        # Create the initial form
                        result = form_service.forms().create(body=form).execute()

                        # Prepare the update requests to include the multiple-choice questions
                        update_requests = [
                            {
                                "updateSettings": {
                                    "settings": {"quizSettings": {"isQuiz": True}},
                                    "updateMask": "quizSettings.isQuiz",
                                }
                            },
                            {
                                "updateFormInfo": {
                                    "info": {
                                        "description": "Please complete this quiz based on this material."
                                    },
                                    "updateMask": "description",
                                }
                            }
                        ]

                        question_index = 0

                        # Parse the questions and add them to the form
                        for question_data in forms_questions:
                            question_text = question_data.get("question", "")
                            options = [{"value": option} for option in question_data.get("options", [])]
                            correct_answer = question_data.get("correctAnswer", "A")

                            # Map correctAnswer letter to index
                            parsed_correct_answer = ord(correct_answer.lower()) - ord('a')

                            # Prepare the request for creating a question
                            question_request = {
                                "createItem": {
                                    "item": {
                                        "title": question_text,
                                        "questionItem": {
                                            "question": {
                                                "required": True,
                                                "grading": {
                                                    "pointValue": 1,
                                                    "correctAnswers": {
                                                        "answers": [{"value": options[parsed_correct_answer]["value"]}]
                                                    },
                                                    "whenRight": {"text": "You got it!"},
                                                    "whenWrong": {"text": "Sorry, that's wrong"}
                                                },
                                                "choiceQuestion": {
                                                    "type": "RADIO",
                                                    "options": options,
                                                    "shuffle": True
                                                }
                                            }
                                        }
                                    },
                                    "location": {"index": question_index}
                                }
                            }

                            # Add the request to the list and increment the index
                            update_requests.append(question_request)
                            question_index += 1  # Increment index for the next question

                        # Now batch update with all questions
                        try:
                            form_service.forms().batchUpdate(formId=result["formId"], body={"requests": update_requests}).execute()
                            form_url = f"https://docs.google.com/forms/d/{result['formId']}/edit"
                            webbrowser.open(form_url)

                            form_txt = io.BytesIO()
                            form_txt.write(form_url.encode('utf-8'))  # Write the form URL to BytesIO
                            form_txt.seek(0)
                            # Save the form URL in a text file
                            # form_file_path = os.path.join(os.getcwd(), 'form.txt')
                            # with open(form_file_path, 'w') as form_file:
                            #     form_file.write(form_url)
                            print("Form link saved to form.txt")

                        except Exception as e:
                            print(f"Failed to update form with questions: {str(e)}")

                    except json.JSONDecodeError as e:
                        print(f"Error decoding JSON: {str(e)}")
                        print(f"Form questions generated but could not parse the response.")

                    return {
                        'response.docx': send_file(response_doc, as_attachment=True, download_name='response.docx'),
                        'questions.docx': send_file(questions_doc, as_attachment=True, download_name='questions.docx'),
                        'answers.docx': send_file(answers_doc, as_attachment=True, download_name='answers.docx'),
                        'form.txt': send_file(form_txt, as_attachment=True, download_name='form.txt')
                    }

                else:
                    print("No response received from the AI.")
            else:
                print("Answers not generated")
        else:
            print("Questions not generated")
    else:
        print("Notes not generated")
