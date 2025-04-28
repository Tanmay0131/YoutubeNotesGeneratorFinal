import re
import io
from youtube_transcript_api import YouTubeTranscriptApi
import google.generativeai as genai
from flask import Flask, request, send_file, render_template
from docx import Document
import json
import webbrowser
from apiclient import discovery
from oauth2client.file import Storage
from httplib2 import Http
from oauth2client import tools, client
import uuid
from pytubefix import YouTube
from pytubefix.cli import on_progress
import moviepy.editor as mp
import speech_recognition as sr
from pydub import AudioSegment
from pydub.silence import split_on_silence
from pathlib import Path
import os

app = Flask(__name__)

generated_files = {}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_text():
    try:
        youtube_link = request.json.get('youtubeLink')
        api_key = request.json.get('apiKey')

        if not youtube_link:
            return 'YouTube link is required', 400
        if not api_key:
            return 'API key is required', 400

        genai.configure(api_key=api_key)

        # try:
        transcript = YouTubeTranscriptApi.get_transcript(youtube_link.replace("https://www.youtube.com/watch?v=", "").strip(), languages=['en'])
        combined_text = ' '.join([item['text'] for item in transcript])
        # except:
        #     transcript = get_transcript_from_youtube_url(youtube_link)
        #     combined_text = transcript

        if "https://www.youtube.com/watch?v=" in youtube_link:
            youtube_link = youtube_link.replace("https://www.youtube.com/watch?v=", "").strip()

        # Generate notes
        prompt = f"""
        Make clean and organized notes using the following transcript. Make sure the information is relevant and concise.

        Transcript: "{combined_text}"
        """
        response = genai.GenerativeModel('gemini-1.5-flash').generate_content(prompt)

        if hasattr(response, 'text'):
            cleaned_text = response.text.replace('*', '')

            response_doc = io.BytesIO()
            document = Document()
            document.add_heading('Generated Notes', level=1)
            document.add_paragraph('Taken from: ' + youtube_link)
            document.add_paragraph(cleaned_text)
            document.save(response_doc)
            response_doc.seek(0)
            generated_files['response'] = response_doc

            # Generate questions
            question_prompt = f"""
            Based on the following notes and transcript, generate a list of questions without answers that would help someone study the material.

            Notes: "{cleaned_text}"
            Transcript: "{combined_text}"
            """
            question_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(question_prompt)

            if hasattr(question_response, 'text'):
                questions_text = question_response.text.replace('*', '')

                questions_doc = io.BytesIO()
                questions_document = Document()
                questions_document.add_heading('Generated Questions', level=1)
                questions_document.add_paragraph('Taken from: ' + youtube_link)
                questions_document.add_paragraph(questions_text)
                questions_document.save(questions_doc)
                questions_doc.seek(0)
                generated_files['questions'] = questions_doc

                # Generate answers
                answers_prompt = f"""
                Provide answers to the following questions based on the notes and transcript.

                Questions: "{questions_text}"
                Notes: "{cleaned_text}"
                Transcript: "{combined_text}"
                """
                answers_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(answers_prompt)

                if hasattr(answers_response, 'text'):
                    answers_text = answers_response.text.replace('*', '')

                    print("Notes: ", cleaned_text)
                    print("Questions: ", questions_text)
                    print("Answers: ", answers_text)

                    answers_doc = io.BytesIO()
                    answers_document = Document()
                    answers_document.add_heading('Generated Answers', level=1)
                    answers_document.add_paragraph('Taken from: ' + youtube_link)
                    answers_document.add_paragraph(answers_text)
                    answers_document.save(answers_doc)
                    answers_doc.seek(0)
                    generated_files['answers'] = answers_doc

                    # Generate Google Forms
                    form_prompt = f"""
                    Based on the following notes and transcript, generate a list of multiple-choice questions. 
                    Format each question as JSON objects with the following structure:

                    {{
                        "question": "Question text",
                        "options": [
                            "Option A",
                            "Option B",
                            "Option C",
                            "Option D"
                        ],
                        "correctAnswer": "A"
                    }}

                    Questions: "{questions_text}"
                    Notes: "{cleaned_text}"
                    Transcript: "{combined_text}"
                    """
                    form_question_response = genai.GenerativeModel('gemini-1.5-flash').generate_content(form_prompt)

                    if hasattr(form_question_response, 'text'):
                        forms_questions_json = form_question_response.text.strip()
                        forms_questions_json = forms_questions_json[7:-3].strip()


                        print("Generated JSON response: ", forms_questions_json)

                        if forms_questions_json:
                            try:
                                forms_questions = json.loads(forms_questions_json)
                                print("Parsed JSON successfully")

                                # Setup Google Forms API
                                SCOPES = [
                                    "https://www.googleapis.com/auth/forms.body",
                                    "https://www.googleapis.com/auth/forms.responses.readonly"
                                ]
                                DISCOVERY_DOC = "https://forms.googleapis.com/$discovery/rest?version=v1"

                                store = Storage("token.json")
                                creds = store.get()
                                if not creds or creds.invalid:
                                    flow = client.flow_from_clientsecrets("client_secrets.json", SCOPES)
                                    creds = tools.run_flow(flow, store)

                                form_service = discovery.build(
                                    "forms", "v1", http=creds.authorize(Http()), discoveryServiceUrl=DISCOVERY_DOC, static_discovery=False
                                )

                                form = {
                                    "info": {
                                        "title": "Generated Quiz",
                                    }
                                }
                                result = form_service.forms().create(body=form).execute()

                                update_requests = [
                                    {
                                        "updateSettings": {
                                            "settings": {"quizSettings": {"isQuiz": True}},
                                            "updateMask": "quizSettings.isQuiz",
                                        }
                                    }
                                ]
                                question_index = 0

                                for question_data in forms_questions:
                                    question_text = question_data.get("question", "")
                                    options = [{"value": option} for option in question_data.get("options", [])]
                                    correct_answer = question_data.get("correctAnswer", "A")

                                    parsed_correct_answer = ord(correct_answer.lower()) - ord('a')

                                    question_request = {
                                        "createItem": {
                                            "item": {
                                                "title": question_text,
                                                "questionItem": {
                                                    "question": {
                                                        "required": True,
                                                        "grading": {
                                                            "pointValue": 1,
                                                            "correctAnswers": {
                                                                "answers": [{"value": options[parsed_correct_answer]["value"]}]
                                                            },
                                                            "whenRight": {"text": "You got it!"},
                                                            "whenWrong": {"text": "Sorry, that's wrong"}
                                                        },
                                                        "choiceQuestion": {
                                                            "type": "RADIO",
                                                            "options": options,
                                                            "shuffle": True
                                                        }
                                                    }
                                                }
                                            },
                                            "location": {
                                                "index": question_index 
                                            }
                                        }
                                    }
                                    update_requests.append(question_request)
                                    question_index += 1


                                form_service.forms().batchUpdate(formId=result["formId"], body={"requests": update_requests}).execute()
                                form_url = f"https://docs.google.com/forms/d/{result['formId']}/edit"
                                webbrowser.open(form_url)

                                form_txt = io.BytesIO()
                                form_txt.write(form_url.encode('utf-8'))
                                form_txt.seek(0)
                                generated_files['form'] = form_txt

                            except json.JSONDecodeError as e:
                                print(f"Error decoding JSON: {str(e)}")
                                return "Invalid JSON format received from the AI", 500
                        else:
                            print("Empty response received for form questions")
                            return "No valid content generated for form questions", 500
                    else:
                        print("No response received from the AI.")
                        return "AI failed to generate form questions", 500

        return json.dumps({"message": "Process completed successfully", "formUrl": form_url})
    except Exception as e:
        print(f"Error: {str(e)}")
        return json.dumps({"error": "Failed to process. Please try again later." }), 500
    

# def download_YouTube_mp4(video_url):
#     yt = YouTube(video_url, on_progress_callback = on_progress)
#     video_stream = yt.streams.get_highest_resolution()
#     video_name = yt.title
#     file_path = video_stream.download()
#     return video_name, file_path


# def convert_mp4_to_wav(file_path, output_path):
#     filename = os.path.splitext(os.path.basename(file_path))[0]
#     # mp.VideoFileClip.ffmpeg_binary = r"C:\\Users\\CMP_OwDiBacco\\Downloads\\Convert Youtube URL to Transcript\\ffmpeg-7.1-full_build (1)\\ffmpeg-7.1-full_build\\bin\\ffmpeg.exe" # replace with path to ffmpeg.exe after you set it as a path in your enviormental varaibles
#     # ffmpeg is not needed ^
#     video = mp.VideoFileClip(file_path)
#     output_wav_dir = os.path.join(output_path, "Wav")
#     os.makedirs(output_wav_dir, exist_ok=True)
#     output_wav_path = os.path.join(output_wav_dir, filename + '.wav')
#     video.audio.write_audiofile(output_wav_path)
#     return output_wav_path


# def convert_wav_to_text(file_path, output_path):
#     chunk_text = []
#     audio = AudioSegment.from_wav(file_path)
#     filename = os.path.splitext(os.path.basename(file_path))[0]
#     chunks = split_on_silence(audio, min_silence_len=500, silence_thresh=-50)
#     if not chunks:
#         return ""  

#     recognizer = sr.Recognizer()
#     output_path = os.path.join(output_path, "Clips", filename)
#     os.makedirs(output_path, exist_ok=True)

#     for i, chunk in enumerate(chunks):
#         full_wav_path = os.path.join(output_path, f"chunk{i+1}.wav")
#         chunk.export(full_wav_path, format="wav")
        
#         with sr.AudioFile(full_wav_path) as source:
#             audio_chunk = recognizer.record(source, duration=4)
#             try:
#                 text = recognizer.recognize_google(audio_chunk)
#                 chunk_text.append(text)
#             except sr.UnknownValueError:
#                 print(f"Chunk {i+1}: No Speech Recognized")
#             except sr.RequestError as e:
#                 print(f"Error With Google Speech Recognition API: {e}")
#             except Exception as e:
#                 print(f"An Error Occurred: {e}")

#     final_text = ''.join(chunk_text)
#     return final_text

# def write_transcript_to_file(text, video_name):
#     with open(os.path.join(video_name + ".txt"), "w") as txt_file:
#         txt_file.write(text)


# def delete_individual_variables(arr):
#     for a in arr:
#         del a


# def delete_created_files(workspace, output_path):
#     delete_script_path = os.path.join(workspace, 'Delete.py')
#     with open(delete_script_path, "w") as wrfile:
#         wrfile.write('import shutil\n')
#         wrfile.write(f'shutil.rmtree(r"{output_path}")\n')
#         wrfile.write('import os\n')
#         wrfile.write(f'os.remove(r"{delete_script_path}")\n')

#     os.system(f'python3 "{delete_script_path}"')


# def get_transcript_from_youtube_url(video_url):
#     output_path_id = str(uuid.uuid4())
#     output_path = f'.\\{output_path_id}'
#     current_dir = Path.cwd()
#     delete_dir = os.path.join(current_dir, output_path_id)
#     video_name, file_path = download_YouTube_mp4(video_url)
#     wav_file_path = convert_mp4_to_wav(file_path, output_path)
#     text = convert_wav_to_text(wav_file_path, output_path)
#     delete_individual_variables([file_path, wav_file_path, output_path])
#     delete_created_files(current_dir, delete_dir)
#     return text
    # write_transcript_to_file(text, video_name)
    # delete_created_files(current_dir, delete_dir)

@app.route('/download/<file_type>')
def download_file(file_type):
    if file_type in generated_files:
        file_obj = generated_files[file_type]
        file_extension = 'docx' if file_type != 'form' else 'txt'
        return send_file(file_obj, as_attachment=True, download_name=f'{file_type}.{file_extension}')
    return 'File not found', 404

if __name__ == '__main__':
    app.run(debug=True, port=5501) 